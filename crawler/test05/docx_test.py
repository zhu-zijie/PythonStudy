from docx import Document
from docx.document import Document as Doc
from docx.shared import Pt, Cm

# create a new Document
document = Document()
# 添加大标题
document.add_heading('Python办公自动化', 0)
# 添加段落
p = document.add_paragraph('Python是一种优雅的语言，它可以用于办公自动化。')
run = p.add_run('简单')
run.bold = True
run.font_size = Pt(20)
p.add_run("而且")
run = p.add_run('优雅')
run.font_size = Pt(20)
run.underline = True
p.add_run('。')

# 添加一级标题
document.add_heading('一级标题', level=1)
# 添加带样式的段落
document.add_paragraph('Intense quote', style='Intense Quote')
# 添加无序列表
document.add_paragraph('first item in unordered list', style='List Bullet')
document.add_paragraph('second item in unordered list', style='List Bullet')

# 添加有序列表
document.add_paragraph('first item in ordered list', style='List Number')
document.add_paragraph('second item i ordered list', style='List Number')

# 添加图片
# document.add_picture('123.jpg', width=Cm(10))

# 添加分节符
document.add_section()

records = (
    ('骆昊', '男', '1995-5-5'),
    ('孙美丽', '女', '1992-2-2')
)
# 添加表格
table = document.add_table(rows=1, cols=3)
table.style = 'Dark List'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '姓名'
hdr_cells[1].text = '性别'
hdr_cells[2].text = '出生日期'
for name, sex, birth in records:
    row_cells = table.add_row().cells
    row_cells[0].text = name
    row_cells[1].text = sex
    row_cells[2].text = birth

# 添加分页符
document.add_page_break()

# 保存文档
document.save('test.docx')
